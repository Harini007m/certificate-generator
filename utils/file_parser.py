"""
File Parser Module
Handles parsing of different file formats (Excel, CSV, DOC, DOCX)
to extract student information.
"""

import pandas as pd
import csv
from docx import Document
import os


class FileParser:
    """Parse student data from various file formats."""

    def parse_file(self, file_path):
        """
        Parse student data from uploaded file.
        
        Args:
            file_path: Path to the uploaded file
            
        Returns:
            List of dictionaries containing student information
        """
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension in ['.xlsx', '.xls']:
            return self._parse_excel(file_path)
        elif file_extension == '.csv':
            return self._parse_csv(file_path)
        elif file_extension in ['.doc', '.docx']:
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    def _parse_excel(self, file_path):
        """Parse Excel file."""
        try:
            df = pd.read_excel(file_path)
            return self._normalize_dataframe(df)
        except Exception as e:
            raise Exception(f"Error parsing Excel file: {str(e)}")

    def _parse_csv(self, file_path):
        """Parse CSV file."""
        try:
            df = pd.read_csv(file_path)
            return self._normalize_dataframe(df)
        except Exception as e:
            raise Exception(f"Error parsing CSV file: {str(e)}")

    def _parse_docx(self, file_path):
        """Parse DOCX file - assumes table format."""
        try:
            doc = Document(file_path)
            students = []

            # Try to extract data from tables
            for table in doc.tables:
                # Assume first row is header
                headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
                
                for row in table.rows[1:]:
                    student_data = {}
                    for i, cell in enumerate(row.cells):
                        if i < len(headers):
                            student_data[headers[i]] = cell.text.strip()
                    
                    if student_data:
                        students.append(self._normalize_student_data(student_data))

            # If no tables found, try to parse from paragraphs
            if not students:
                students = self._parse_docx_paragraphs(doc)

            return students
        except Exception as e:
            raise Exception(f"Error parsing DOCX file: {str(e)}")

    def _parse_docx_paragraphs(self, doc):
        """Parse DOCX paragraphs looking for student data."""
        students = []
        current_student = {}

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Try to identify key-value pairs
                if ':' in text:
                    parts = text.split(':', 1)
                    key = parts[0].strip().lower()
                    value = parts[1].strip()
                    current_student[key] = value
                elif current_student:
                    # Start new student
                    students.append(self._normalize_student_data(current_student))
                    current_student = {}

        if current_student:
            students.append(self._normalize_student_data(current_student))

        return students

    def _normalize_dataframe(self, df):
        """Convert dataframe to list of normalized student dictionaries."""
        students = []
        
        # Convert column names to lowercase for easier matching
        df.columns = df.columns.str.lower().str.strip()

        for _, row in df.iterrows():
            student_data = row.to_dict()
            normalized = self._normalize_student_data(student_data)
            if normalized:
                students.append(normalized)

        return students

    def _normalize_student_data(self, data):
        """
        Normalize student data to standard format.
        Expected fields: name, department, class, email
        """
        normalized = {}

        # Map various field names to standard names
        name_fields = ['name', 'student_name', 'student name', 'full_name', 'full name', 'studentname']
        dept_fields = ['department', 'dept', 'branch', 'stream', 'course']
        class_fields = ['class', 'year', 'semester', 'grade', 'level', 'section']
        email_fields = ['email', 'email_address', 'email address', 'mail', 'e-mail']

        # Extract name
        for field in name_fields:
            if field in data and data[field]:
                normalized['name'] = str(data[field]).strip()
                break

        # Extract department
        for field in dept_fields:
            if field in data and data[field]:
                normalized['department'] = str(data[field]).strip()
                break

        # Extract class
        for field in class_fields:
            if field in data and data[field]:
                normalized['class'] = str(data[field]).strip()
                break

        # Extract email
        for field in email_fields:
            if field in data and data[field]:
                normalized['email'] = str(data[field]).strip()
                break

        # Only return if at least name is present
        if 'name' in normalized:
            # Set defaults for optional fields
            normalized.setdefault('department', 'N/A')
            normalized.setdefault('class', 'N/A')
            normalized.setdefault('email', '')
            return normalized

        return None
